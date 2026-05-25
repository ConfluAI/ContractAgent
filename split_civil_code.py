"""
Split 中华人民共和国民法典 .docx into chunks.
Extracts:
  - Book 1 (总则): Chapters 6 (民事法律行为), 7 (代理), 8 (民事责任)
  - Book 3 (合同): Entire book

Output JSONL: {page_content, metadata}
page_content embeds hierarchical context as a readable prefix.
"""

import re
import json
import sys
from pathlib import Path
from docx import Document

# ── Chinese numeral conversion ──────────────────────────────────────────

CN_NUM_MAP = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9,
    "十": 10, "百": 100, "千": 1000, "万": 10000,
    "零": 0,
}


def cn_to_int(cn: str) -> int | None:
    if not cn:
        return None
    result = 0
    current = 0
    for ch in cn:
        val = CN_NUM_MAP.get(ch)
        if val is None:
            return None
        if ch in ("十", "百", "千", "万"):
            if current == 0:
                current = 1
            result += current * val
            current = 0
        else:
            current = val
    result += current
    return result


# ── Heading / article detection ─────────────────────────────────────────

HEADING_RE = re.compile(
    r"^第([一二三四五六七八九十百千万零]+)"
    r"(编|分编|章|节)"
    r"[　\s]+(.+)"
)

ARTICLE_RE = re.compile(
    r"^第([一二三四五六七八九十百千万零]+)条[　\s]"
)

FU_ZE_RE = re.compile(r"^附[　\s]+则$")


def parse_heading(text: str) -> dict | None:
    if FU_ZE_RE.match(text):
        return {"level": "supplementary", "number": None, "title": "附则",
                "prefix": "附则"}

    m = HEADING_RE.match(text)
    if not m:
        return None

    cn_num, level_name, title = m.groups()
    num = cn_to_int(cn_num)
    level_map = {
        "编": "book", "分编": "sub_book", "章": "chapter", "节": "section",
    }
    title_clean = re.sub(r"[　]+", "", title.strip())
    return {
        "level": level_map.get(level_name, level_name),
        "number": num,
        "title": title_clean,
        "prefix": f"第{cn_num}{level_name}",
    }


def is_article_start(text: str) -> tuple[str, int] | None:
    m = ARTICLE_RE.match(text)
    if not m:
        return None
    cn = m.group(1)
    num = cn_to_int(cn)
    return (cn, num) if num is not None else None


# ── Content extraction ──────────────────────────────────────────────────

BOOK1_INCLUDE_CHAPTERS = {6, 7, 8}  # 民事法律行为, 代理, 民事责任


def should_include(book_num: int | None, chapter_num: int | None) -> bool:
    if book_num == 3:
        return True
    if book_num == 1 and chapter_num in BOOK1_INCLUDE_CHAPTERS:
        return True
    return False


def extract_articles(filepath: str) -> list[dict]:
    doc = Document(filepath)

    # Find first article (content start)
    content_start = 0
    for i, p in enumerate(doc.paragraphs):
        if is_article_start(p.text.strip()):
            content_start = i
            break

    # Re-track to capture headings right before first article
    heading_start = content_start
    while heading_start > 0:
        prev = doc.paragraphs[heading_start - 1].text.strip()
        if parse_heading(prev) or FU_ZE_RE.match(prev) or prev == "":
            heading_start -= 1
        else:
            break

    results = []
    ctx: dict = {
        "book": None, "book_num": None, "book_prefix": None,
        "sub_book": None, "sub_book_prefix": None,
        "chapter": None, "chapter_num": None, "chapter_prefix": None,
        "section": None, "section_prefix": None,
    }
    current_article: dict | None = None

    for i in range(heading_start, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue

        heading = parse_heading(text)
        if heading:
            lvl = heading["level"]
            if lvl == "book":
                ctx.update(book=heading["title"], book_num=heading["number"],
                           book_prefix=heading["prefix"])
                ctx.update(sub_book=None, sub_book_prefix=None, chapter=None,
                           chapter_num=None, chapter_prefix=None, section=None,
                           section_prefix=None)
            elif lvl == "sub_book":
                ctx.update(sub_book=heading["title"],
                           sub_book_prefix=heading["prefix"])
                ctx.update(chapter=None, chapter_num=None, chapter_prefix=None,
                           section=None, section_prefix=None)
            elif lvl == "chapter":
                ctx.update(chapter=heading["title"], chapter_num=heading["number"],
                           chapter_prefix=heading["prefix"])
                ctx.update(section=None, section_prefix=None)
            elif lvl == "section":
                ctx.update(section=heading["title"],
                           section_prefix=heading["prefix"])
            elif lvl == "supplementary":
                ctx.update(chapter="附则", chapter_num=None, chapter_prefix="附则",
                           section=None, section_prefix=None)
            continue

        article_match = is_article_start(text)
        if article_match:
            if current_article and should_include(ctx["book_num"], ctx["chapter_num"]):
                results.append(_build_chunk(current_article, ctx))
            cn, num = article_match
            body = ARTICLE_RE.sub("", text).strip()
            current_article = {"cn": cn, "num": num, "lines": [body]}
            continue

        if i < content_start:
            continue

        if current_article is not None:
            current_article["lines"].append(text)

    if current_article and should_include(ctx["book_num"], ctx["chapter_num"]):
        results.append(_build_chunk(current_article, ctx))

    return results


def _build_chunk(article: dict, ctx: dict) -> dict:
    body = "".join(article["lines"])

    # Build page_content with hierarchical prefix
    source_line = "中华人民共和国民法典"
    book_line = f"{ctx['book_prefix']} {ctx['book']}" if ctx.get("book_prefix") else ""
    sub_book_line = f"{ctx['sub_book_prefix']} {ctx['sub_book']}" if ctx.get("sub_book_prefix") else ""
    chapter_line = f"{ctx['chapter_prefix']} {ctx['chapter']}" if ctx.get("chapter_prefix") else ""
    section_line = f"{ctx['section_prefix']} {ctx['section']}" if ctx.get("section_prefix") else ""
    article_line = f"第{article['cn']}条 {body}"

    lines = [source_line]
    if book_line:
        lines.append(book_line)
    if sub_book_line:
        lines.append(sub_book_line)
    if chapter_line:
        lines.append(chapter_line)
    if section_line:
        lines.append(section_line)
    lines.append(article_line)

    page_content = "\n".join(lines)

    # Build structured metadata
    meta = {
        "source": "中华人民共和国民法典",
        "section_id": f"第{article['cn']}条",
        "article_num": article["num"],
        "book": book_line,
        "chapter": chapter_line,
        "law_rank": 4,
        "law_rank_desc": "法律",
    }
    if sub_book_line:
        meta["sub_book"] = sub_book_line
    if section_line:
        meta["section"] = section_line
    meta["domain"] = _classify_domain(meta)

    return {"page_content": page_content, "metadata": meta}


def _classify_domain(meta: dict) -> str:
    chapter = meta.get("chapter", "")
    book = meta.get("book", "")

    if "合同" in book:
        if any(t in chapter for t in [
            "买卖", "借款", "租赁", "承揽", "建设工程", "运输", "技术",
            "保管", "仓储", "委托", "物业", "行纪", "中介", "合伙", "赠与",
            "融资租赁", "保理", "供用电",
        ]):
            return "合同分则"
        if "保证" in chapter:
            return "保证合同"
        if "一般规定" in chapter:
            return "合同通则"
        if "订立" in chapter:
            return "合同订立"
        if "效力" in chapter:
            return "合同效力"
        if "履行" in chapter:
            return "合同履行"
        if "保全" in chapter:
            return "合同保全"
        if "变更" in chapter or "转让" in chapter:
            return "合同变更与转让"
        if "终止" in chapter:
            return "合同终止"
        if "违约" in chapter:
            return "违约责任"
        return "合同通用"
    if "代理" in chapter:
        return "代理"
    if "民事法律行为" in chapter:
        return "民事法律行为"
    if "责任" in chapter:
        return "民事责任"
    return "合同通用"


# ── Main ────────────────────────────────────────────────────────────────

def main():
    input_path = Path("data/中华人民共和国民法典_20200528.docx")
    output_path = Path("data/civil_code_contract_chunks.jsonl")

    print(f"Reading {input_path}...", file=sys.stderr)
    chunks = extract_articles(str(input_path))
    print(f"Extracted {len(chunks)} articles.", file=sys.stderr)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        for chunk in chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    print(f"Done. Output: {output_path}", file=sys.stderr)

    # Sample
    if chunks:
        print("\n── Sample ──", file=sys.stderr)
        print(chunks[0]["page_content"][:300], file=sys.stderr)
        print("...", file=sys.stderr)
        print(json.dumps(chunks[0]["metadata"], ensure_ascii=False, indent=2),
              file=sys.stderr)


if __name__ == "__main__":
    main()
